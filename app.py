# -*- coding: utf-8 -*-
"""
Created on Fri Aug 11 12:40:42 2023

@author: AFOMIKE
"""
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, flash, redirect, url_for
import os
from numpy import vectorize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import secrets

app = Flask(__name__)
# app.secret_key = 'your_secret_key_here'

sample_files = []
sample_contents = []
s_vectors = []

# def check_plagiarism():
#     results = set()
#     for sample_a, text_vector_a in s_vectors:
#         new_vectors = s_vectors.copy()
#         current_index = new_vectors.index((sample_a, text_vector_a))
#         del new_vectors[current_index]
#         for sample_b, text_vector_b in new_vectors:
#             sim_score = similarity(text_vector_a, text_vector_b)[0][1]
#             sample_pair = sorted((sample_a, sample_b))
#             score = sample_pair[0], sample_pair[1], sim_score
#             results.add(score)
#     return results
def check_plagiarism():
    results = set()
    for sample_a, text_vector_a in s_vectors:
        new_vectors = s_vectors.copy()
        current_index = new_vectors.index((sample_a, text_vector_a))
        del new_vectors[current_index]
        for sample_b, text_vector_b in new_vectors:
            sim_score = similarity(text_vector_a, text_vector_b)[0][1]
            similarity_percentage = sim_score * 100  # Convert to percentage
            similarity_percentage =  f"{similarity_percentage:.1f}%" # Convert to percentage
            sample_pair = sorted((sample_a, sample_b))
            score = sample_pair[0], sample_pair[1], similarity_percentage
            results.add(score)
    return results

# def vectorize_samples():
#     global sample_files, sample_contents, s_vectors
#     upload_folder = 'mysite/uploads'

#     # Get a list of files with their modification times
#     files_with_times = [(file, os.path.getmtime(os.path.join(upload_folder, file))) for file in os.listdir(upload_folder) if file.endswith(('.txt', '.docx', '.doc'))]

#     # Sort files based on modification time in descending order (newest first)
#     files_with_times.sort(key=lambda x: x[1], reverse=True)

#     # Extract sorted file names
#     sample_files = [file[0] for file in files_with_times]

#     sample_contents = []
#     for file in sample_files:
#         file_path = os.path.join(upload_folder, file)
#         content_type, _ = mimetypes.guess_type(file_path)

#         print(f"File: {file}, Content Type: {content_type}")

#         if content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
#             try:
#                 if file.endswith('.txt'):
#                     with open(file_path, 'r') as f:
#                         content = f.read()
#                         sample_contents.append(content)
#                 elif file.endswith(('.docx', '.doc')):
#                     doc = Document(file_path)
#                     content = ' '.join([para.text for para in doc.paragraphs])
#                     sample_contents.append(content)
#             except ValueError as e:
#                 print(f"Error reading {file}: {e}")
#         else:
#             print(f"Skipped {file} due to unsupported content type: {content_type}")

#     s_vectors = list(zip(sample_files, vectorize(sample_contents)))

# def vectorize_samples():
#     global sample_files, sample_contents, s_vectors
#     sample_files = [doc for doc in os.listdir('uploads') if doc.endswith(('.txt', '.docx', '.doc'))]
#     sample_contents = []
#     for file in sample_files:
#         if file.endswith('.txt'):
#             with open(os.path.join('uploads', file), 'r') as f:
#                 content = f.read()
#                 sample_contents.append(content)
#         elif file.endswith(('.docx', '.doc')):
#             doc = Document(os.path.join('uploads', file))
#             content = ' '.join([para.text for para in doc.paragraphs])
#             sample_contents.append(content)
#     s_vectors = list(zip(sample_files, vectorize(sample_contents)))
# vectorize = lambda Text: TfidfVectorizer().fit_transform(Text).toarray()
# similarity = lambda doc1, doc2: cosine_similarity([doc1, doc2])

def vectorize_samples():
    global sample_files, sample_contents, s_vectors
    upload_folder = 'uploads'
    
    # Get a list of files with their modification times
    files_with_times = [(file, os.path.getmtime(os.path.join(upload_folder, file))) for file in os.listdir(upload_folder) if file.endswith(('.txt', '.docx', '.doc'))]
    
    # Sort files based on modification time in descending order (newest first)
    files_with_times.sort(key=lambda x: x[1], reverse=True)
    
    # Extract sorted file names
    sample_files = [file[0] for file in files_with_times]
    
    sample_contents = []
    for file in sample_files:
        if file.endswith('.txt'):
            with open(os.path.join(upload_folder, file), 'r') as f:
                content = f.read()
                sample_contents.append(content)
        elif file.endswith(('.docx', '.doc')):
            doc = Document(os.path.join(upload_folder, file))
            content = ' '.join([para.text for para in doc.paragraphs])
            sample_contents.append(content)
    
    s_vectors = list(zip(sample_files, vectorize(sample_contents)))

vectorize = lambda Text: TfidfVectorizer().fit_transform(Text).toarray()
similarity = lambda doc1, doc2: cosine_similarity([doc1, doc2])
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    uploaded_file = request.files['file']
    if uploaded_file.filename != '':
        # Ensure the filename is safe for storage
        filename = secure_filename(uploaded_file.filename)
        file_path = os.path.join('uploads', filename)
        uploaded_file.save(file_path)
        flash('File uploaded successfully', 'success')
        vectorize_samples()
        return redirect(url_for('index'))
    flash('No file selected for upload', 'error')
    return redirect(url_for('index'))

@app.route('/plagiarism')
def plagiarism():
    results = check_plagiarism()
    return render_template('index.html', results=results)

if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    key = secrets.token_hex(16)  # Generates a 32-character random hexadecimal string
    app.secret_key = key
    app.run(debug=True)
